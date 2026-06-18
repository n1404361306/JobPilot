import uuid
from pathlib import Path

from fastapi import UploadFile

from app.core.config import get_settings
from app.core.exceptions import BusinessException
from app.modules.ocr.service import OCRService

RESUME_UPLOAD_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".png", ".jpg", ".jpeg", ".webp"}
RESUME_UPLOAD_MIME_TYPES = {
    "application/pdf",
    "application/x-pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "text/plain",
    "text/markdown",
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/webp",
    "application/octet-stream",
    "binary/octet-stream",
}


async def save_resume_upload(*, user_id: int, upload: UploadFile) -> tuple[str, str, str, int]:
    settings = get_settings()
    original_name = upload.filename or "resume.pdf"
    suffix = Path(original_name).suffix.lower()

    if suffix not in RESUME_UPLOAD_EXTENSIONS:
        raise BusinessException(
            code=4006,
            message="不支持的文件类型，请上传 PDF、DOCX、TXT、MD、PNG、JPG 或 WEBP",
        )

    content_type = (upload.content_type or "").lower()
    if content_type and content_type not in RESUME_UPLOAD_MIME_TYPES:
        raise BusinessException(code=4006, message=f"不支持的文件 MIME 类型：{content_type}")

    data = await upload.read()
    max_bytes = settings.ocr_max_file_size_mb * 1024 * 1024
    if len(data) > max_bytes:
        raise BusinessException(code=4007, message="文件过大，请压缩后重试")

    base_dir = Path(settings.upload_dir) / "resume-uploads" / str(user_id)
    base_dir.mkdir(parents=True, exist_ok=True)

    stored_name = f"{uuid.uuid4().hex}{suffix}"
    file_path = base_dir / stored_name
    file_path.write_bytes(data)

    return str(file_path), original_name, content_type or suffix, len(data)


def extract_docx_text(file_path: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise BusinessException(code=5005, message="服务端未安装 python-docx，请联系管理员") from exc

    document = Document(file_path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise BusinessException(code=4008, message="Word 文档中未提取到文字内容")
    return text


def extract_plain_text(file_path: str) -> str:
    data = Path(file_path).read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            text = data.decode(encoding).strip()
            break
        except UnicodeDecodeError:
            continue
    else:
        raise BusinessException(code=4008, message="文本文件编码不支持，请保存为 UTF-8 后重试")

    if not text:
        raise BusinessException(code=4008, message="文本文件中未提取到文字内容")
    return text


def extract_resume_text(file_path: str) -> tuple[str, str, str]:
    suffix = Path(file_path).suffix.lower()
    ocr = OCRService()

    try:
        if suffix == ".pdf":
            result = ocr.extract_from_pdf_path(file_path)
            if result.engine in {"pymupdf-text", "pymupdf-text-partial"}:
                source_type = "pdf-text"
            else:
                source_type = "pdf-ocr"
            return result.text.strip(), source_type, result.engine

        if suffix == ".docx":
            return extract_docx_text(file_path), "docx", "python-docx"

        if suffix in {".txt", ".md"}:
            return extract_plain_text(file_path), "plain-text", "text"

        if suffix in {".png", ".jpg", ".jpeg", ".webp"}:
            result = ocr.extract_from_image_path(file_path)
            return result.text.strip(), "image-ocr", result.engine
    except BusinessException:
        raise
    except Exception as exc:
        raise BusinessException(code=5005, message=f"文件解析失败：{exc}") from exc

    raise BusinessException(code=4006, message="不支持的文件类型")
